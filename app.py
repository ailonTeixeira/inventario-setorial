from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file
from docx import Document
from docx.shared import Inches, Pt
from functools import wraps
import database as db
import io
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font

app = Flask(__name__)
app.secret_key = 'chave-secreta-muito-segura-para-a-sessao'

# --- Autenticação ---
USUARIO_VALIDO = "admin"
SENHA_VALIDA = "admin123"

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'logged_in' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# Disponibiliza a lista de setores para todos os templates
@app.context_processor
def inject_setores():
    if 'logged_in' in session:
        return dict(todos_setores=db.get_todos_setores())
    return dict(todos_setores=[])

# --- Rotas ---
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        if request.form['username'] == USUARIO_VALIDO and request.form['password'] == SENHA_VALIDA:
            session['logged_in'] = True
            return redirect(url_for('home'))
        else:
            flash('Usuário ou senha inválidos!', 'danger')
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.pop('logged_in', None)
    return redirect(url_for('login'))

@app.route('/', methods=['GET', 'POST'])
@login_required
def home():
    if request.method == 'POST':
        nome_setor = request.form['nome_setor'].strip()
        if not nome_setor:
            flash('O nome do setor não pode ser vazio.', 'warning')
            return redirect(url_for('home'))

        setor = db.get_setor_por_nome(nome_setor)
        if setor:
            # Se o setor existe, vai para a página dele
            return redirect(url_for('setor_detalhe', setor_id=setor['id']))
        else:
            # Se não existe, cria um novo e vai para a página dele
            novo_setor_id = db.adicionar_setor(nome_setor)
            flash(f'Setor "{nome_setor}" criado com sucesso!', 'success')
            return redirect(url_for('setor_detalhe', setor_id=novo_setor_id))

    setores = db.get_todos_setores()
    return render_template('home.html', setores=setores)

@app.route('/setor/<int:setor_id>')
@login_required
def setor_detalhe(setor_id):
    setor = db.get_setor_por_id(setor_id)
    materiais = db.get_materiais_por_setor(setor_id)
    return render_template('setor_detalhe.html', setor=setor, materiais=materiais)

@app.route('/material/adicionar', methods=['POST'])
@login_required
def adicionar_material_route():
    dados = {
        'setor_id': request.form['setor_id'],
        'nome': request.form['nome'],
        'marca': request.form['marca'],
        'tombo': request.form['tombo'],
        'estado': request.form['estado']
    }
    db.adicionar_material(dados)
    flash('Material adicionado com sucesso!', 'success')
    return redirect(url_for('setor_detalhe', setor_id=dados['setor_id']))

@app.route('/material/editar/<int:material_id>', methods=['POST'])
@login_required
def editar_material_route(material_id):
    dados = {
        'nome': request.form['nome'],
        'marca': request.form['marca'],
        'tombo': request.form['tombo'],
        'estado': request.form['estado']
    }
    db.atualizar_material(material_id, dados)
    flash('Material atualizado com sucesso!', 'success')
    setor_id = request.form['setor_id']
    return redirect(url_for('setor_detalhe', setor_id=setor_id))

@app.route('/material/apagar/<int:material_id>', methods=['POST'])
@login_required
def apagar_material_route(material_id):
    setor_id = request.form['setor_id']
    db.apagar_material(material_id)
    flash('Material apagado com sucesso!', 'warning')
    return redirect(url_for('setor_detalhe', setor_id=setor_id))

# Em app.py

# ... (outras rotas) ...
# Em app.py

@app.route('/setor/<int:setor_id>/relatorio')
@login_required
def gerar_relatorio_word(setor_id):
    # 1. Obter os dados (a consulta já ordena os materiais por nome)
    setor = db.get_setor_por_id(setor_id)
    materiais = db.get_materiais_por_setor(setor_id)

    document = Document('template_relatorio.docx')

    # 2. Adicionar o título
    p = document.add_paragraph()
    run = p.add_run(f'Setor: {setor["nome"]}')
    run.bold = True
    run.font.size = Pt(16)

    # 3. Adicionar a tabela com os dados e numeração
    table = document.add_table(rows=1, cols=5) # <--- MUDANÇA: 5 colunas agora
    table.style = 'Table Grid'

    # Adicionar o cabeçalho da tabela
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Material'
    hdr_cells[2].text = 'Marca'
    hdr_cells[3].text = 'Tombo/Patrimônio'
    hdr_cells[4].text = 'Estado'

    # Adicionar as linhas de dados com os materiais e um contador
    for i, material in enumerate(materiais):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1) # <--- NOVO: Contador do item
        row_cells[1].text = material['nome']
        row_cells[2].text = material['marca'] or ''
        row_cells[3].text = material['tombo'] or ''
        row_cells[4].text = material['estado']

    # 4. Adicionar a frase no final do documento
    document.add_paragraph() # Adiciona um parágrafo em branco para espaçamento
    paragrafo_final = document.add_paragraph('SETOR PATRIMÔNIO')
    paragrafo_final.alignment = WD_ALIGN_PARAGRAPH.CENTER # <--- Centraliza o texto

    # 5. Salvar e enviar o arquivo
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    nome_arquivo_download = f'Relatorio_{setor["nome"].replace(" ", "_")}.docx'
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=nome_arquivo_download,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/relatorio-geral-excel')
@login_required
def gerar_relatorio_geral_excel(): # <--- ESTE É O NOME (ENDPOINT) QUE O FLASK USA
    # 1. Obter todos os setores
    setores = db.get_todos_setores()

    # 2. Criar um novo Workbook (arquivo Excel)
    wb = Workbook()
    ws = wb.active # Pega a primeira planilha
    ws.title = "Inventário Geral"

    # 3. Adicionar o cabeçalho principal
    ws.append(['Setor', 'Item', 'Material', 'Marca', 'Tombo/Patrimônio', 'Estado'])
    # Aplicar negrito no cabeçalho
    for cell in ws[1]:
        cell.font = Font(bold=True)
    
    # 4. Iterar sobre cada setor e seus materiais
    for setor in setores:
        materiais = db.get_materiais_por_setor(setor['id'])
        if not materiais:
            # Se o setor estiver vazio, ainda assim adicionamos uma linha para ele
            ws.append([setor['nome'], '-', 'Nenhum material cadastrado', '-', '-', '-'])
        else:
            for i, material in enumerate(materiais):
                # Na primeira linha do setor, mostramos o nome dele. Nas seguintes, deixamos em branco.
                nome_setor_para_mostrar = setor['nome'] if i == 0 else ''
                ws.append([
                    nome_setor_para_mostrar,
                    i + 1,
                    material['nome'],
                    material['marca'] or '',
                    material['tombo'] or '',
                    material['estado']
                ])
    
    # 5. Salvar o arquivo em memória e enviar para download
    file_stream = io.BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name='Inventario_Geral_Setorial.xlsx',
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5001)
